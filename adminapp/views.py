from django.utils import timezone
from django.http import HttpResponse
from django.shortcuts import get_object_or_404, render,redirect,reverse
from django.contrib import messages
from accounts.models import CustomUser
from category.models import Category, Subcategory,Product,Color,Variations
from orders.models import Order
from django.utils.text import slugify
from django.contrib.auth import authenticate,login,logout
from django.contrib.auth.decorators import login_required
from orders.models import Order,OrderItem
from django.db.models import Sum
import datetime
from django.db.models import Count,Q
import csv
from io import BytesIO
from django.http import HttpResponse
from xhtml2pdf import pisa
from django.views.generic import View
from django.template.loader import get_template
import xlwt
from offers.forms import CategoryOfferForm,SubcategoryOfferForm,ProductOfferForm,CouponForm
from offers.models import CategoryOffer,SubcategoryOffer,ProductOffer,Coupon
from category.models import Color,Variations
from django.core.paginator import Paginator
import io
from docx import Document
from django.http import HttpResponse


def admin_login(request):
    if request.method=='POST':
        email=request.POST['email']
        password=request.POST['pass']
        user=authenticate(email=email,password=password)
    
        if user is not None:
            login(request,user)
            print(request.user.is_admin)
            if request.user.is_admin==True:
                return redirect('admin-home')
            else:
                messages.error(request,'Your Email or Password is Incorrect!!')
                return redirect('adminlogin')
        else:
            messages.error(request,"Your Email or Password is Incorrect!!")
            return redirect('adminlogin')
    return render(request,'adminapp/login.html')



@login_required(login_url='adminlogin')
def admin_home(request):
    if request.user.is_superuser==True:
        total_users=CustomUser.objects.filter(is_active=True).count()
        total_products=Product.objects.filter(is_available=True).count()
        total_orders=OrderItem.objects.filter(status='Delivered').count()
        total_revenue=OrderItem.objects.filter(status='Delivered').aggregate(Sum('price'))
        
        
        
        # Daily sales
        current_year=timezone.now().year
        order_detail=OrderItem.objects.filter(created_at__lt=datetime.date(current_year,12,31),status="Delivered")
        monthly_order_count=[]
        month=timezone.now().month
        for i in range(1,month+2):
            monthly_order = order_detail.filter(created_at__month=i).count()
            monthly_order_count.append(monthly_order)
        
        # Monthly sales
        today=datetime.datetime.now()
        dates=OrderItem.objects.filter(created_at__month=today.month).values('created_at__date').annotate(order_items=Count('id')).order_by('created_at__date')
        returns=OrderItem.objects.filter(created_at__month=today.month).values('created_at__date').annotate(returns=Count('id',filter=Q(status='Cancelled'))).order_by('created_at__date')
        sales=OrderItem.objects.filter(created_at__month=today.month).values('created_at__date').annotate(sales=Count('id',filter=Q(status='Delivered'))).order_by('created_at__date')
        
        # Most moving prduct
        most_moving_product_count=[]
        most_moving_product=[]
        
        products=Product.objects.all()
        for i in products:
            most_moving_product.append(i)
            most_moving_product_count.append(
                OrderItem.objects.filter(
                    product=i,status='Delivered'
                ).count()
            )
        # Status count
        placed_count=OrderItem.objects.filter(status="Order Placed").count()
        shipped_count=OrderItem.objects.filter(status="Shipped").count()
        delivered_count=OrderItem.objects.filter(status="Delivered").count()
        return_count=OrderItem.objects.filter(status="Refund Initiated").count()
        cancelled_count=OrderItem.objects.filter(status="Cancelled").count()
        return render(request,'adminapp/dashboard.html',{
            
            'total_orders':total_orders,
            'total_products':total_products,
            'total_users':total_users,
            'total_revenue':total_revenue,
            "monthly_order_count": monthly_order_count,
            'today':today,
            'sales':sales,
            'returns':returns,
            'dates':dates,
            'most_moving_product':most_moving_product,
            'most_moving_product_count':most_moving_product_count,
            'status_count':[
                placed_count,
                shipped_count,
                delivered_count,
                return_count,
                cancelled_count,
            ]

        })


@login_required(login_url='adminlogin')
def admin_logout(request):
    logout(request)
    return redirect('adminlogin')

# User Managment
@login_required(login_url='adminlogin')
def show_users(request):
    users=CustomUser.objects.all().order_by('-id')
    paginator=Paginator(users,10)
    page=request.GET.get('page')
    paged_user_list=paginator.get_page(page)
    return render(request,'adminapp/users.html',{
        'users':paged_user_list
    })


def block_user(request,id):
    user=CustomUser.objects.get(id=id)
    if user.blocked:
        user.blocked=False
        user.save()
        
    else:
        user.blocked=True
        user.save()
    return redirect('users')


# Category Managment
@login_required(login_url='adminlogin')
def category(request):
    cat=Category.objects.all()
    paginator=Paginator(cat,10)
    page=request.GET.get('page')
    paged_category_list=paginator.get_page(page)
    context={
        'categories':paged_category_list
    }
    return render(request,'adminapp/category.html',context)




@login_required(login_url='adminlogin')
def add_category(request):
    if request.method=="POST":
        if request.POST['catname']:
            name=request.POST['catname']
            if Category.objects.filter(cat_name=name).exists():
                messages.error(request,'Category Already Exists!')
                return redirect('category')
            cat=Category()
            cat.cat_name=name
            cat.slug=slugify(name)
            cat.save()
            messages.error(request,'New Category Added')
            return redirect('category')
        else:
            messages.error(request,'Required all Fields...')
            return redirect('addcategory')

    return render(request,'adminapp/add-category.html')



@login_required(login_url='adminlogin')
def edit_category(request,id):
    cat=Category.objects.get(pk=id)
    if request.method=='POST':
        name=request.POST.get('catname')
        cat.cat_name=name
        cat.slug=slugify(name)
        cat.save()
        messages.error(request,"Updated")
        return redirect('category')
    return render(request,'adminapp/edit-category.html',{
        'cat':cat
    })



@login_required(login_url='adminlogin')
def delete_category(request,id):
    de=Category.objects.filter(id=id).delete()
    return redirect('category')



@login_required(login_url='adminlogin')
def subcategory(request):
    sub_category=Subcategory.objects.all()
    paginator=Paginator(sub_category,10)
    page=request.GET.get('page')
    paged_subcategory_list=paginator.get_page(page)
    context={
        'sub_category':paged_subcategory_list
    }
    return render(request,'adminapp/subcategory.html',context)



@login_required(login_url='adminlogin')
def add_subcategory(request):
    categories = Category.objects.all()

    if request.method == 'POST':
        parent_cat_id = request.POST.get('selectcat')
        sub_name = request.POST.get('subcatname')

        if parent_cat_id and sub_name:
            parent_category = get_object_or_404(Category, id=parent_cat_id)

            # Check if the subcategory with the same name and parent category already exists
            existing_subcategory = Subcategory.objects.filter(sub_name=sub_name, parent_cat=parent_category).first()

            if existing_subcategory:
                # If the subcategory exists, just associate it with the new category.
                messages.success(request, 'Subcategory already exists. Associated with the new category.')
            else:
                # If the subcategory doesn't exist, create a new one.
                subcategory = Subcategory()
                subcategory.sub_name = sub_name
                subcategory.parent_cat = parent_category

                # Generate the slug with parent category's slug to ensure uniqueness.
                subcategory.slug = slugify(sub_name + ' ' + parent_category.slug)

                subcategory.save()
                messages.success(request, 'Subcategory Created')

        else:
            messages.error(request, 'Required all Fields')

        return redirect('addsubcategory')

    return render(request, 'adminapp/add-subcategory.html', {
        'categories': categories
    })


@login_required(login_url='adminlogin')
def edit_subcategory(request, id):
    category = Category.objects.all()
    sub = get_object_or_404(Subcategory, id=id)

    if request.method == 'POST':
        cate_id = request.POST.get('selectcat')
        sub_name = request.POST.get('subcatname')

        if Subcategory.objects.filter(sub_name=sub_name).exclude(id=id).exists():
            messages.error(request, 'Subcategory already exists')
            return redirect('subcategory')
        else:
            category_obj = get_object_or_404(Category, id=cate_id)

            sub.sub_name = sub_name
            sub.slug = slugify(sub_name)
            sub.parent_cat = category_obj
            sub.save()

            messages.success(request, 'Subcategory updated successfully')
            return redirect('subcategory')

    return render(request, 'adminapp/edit-subcategory.html', {
        'categories': category, 'sub': sub
    })

# def edit_subcategory(request,id):
#     category=Category.objects.all()
#     sub=Subcategory.objects.get(id=id)
#     if request.method=='POST':
#         cate=request.POST.get('selectcat')
#         sub_name=request.POST.get('subcatname')
#         if Subcategory.objects.filter(sub_name=sub_name).exists():
#             messages.error(request,'Already Exists')
#             return redirect('subcategory')
#         else:
            
#             sub.sub_name=sub_name
#             sub.slug=slugify(sub_name)
            
#             sub.parent_cat_id=cate
#             sub.save()
#             messages.error(request,'SubCategory Created')
#             return redirect('subcategory')

    
#     return render(request,'adminapp/edit-subcategory.html',{
#         'categories':category,'sub':sub
#     })


@login_required(login_url='adminlogin')
def delete_subcategory(request,id):
    Subcategory.objects.filter(id=id).delete()
    return redirect('subcategory')


@login_required(login_url='adminlogin')
def products(request):
    products_list=Product.objects.all()
    paginator=Paginator(products_list,10)
    page=request.GET.get('page')
    paged_products_list=paginator.get_page(page)
    return render(request,'adminapp/products.html',{
        'products':paged_products_list
    })


@login_required(login_url='adminlogin')
def adding_product(request):
    categories=Category.objects.all()
    subcategories=Subcategory.objects.all()
    context={
        'categories':categories,
        'subcategories':subcategories
    }
    return render(request,'adminapp/addproduct.html',context)



@login_required(login_url='adminlogin')
def load_subcategory(request,catid):
    subcategory=Subcategory.objects.filter(parent_cat_id=catid)
    print(subcategory)
    context={
        'subcategory':subcategory
    }
    return render(request,'adminapp/dropdown.html',context)

@login_required(login_url='adminlogin')
def editload_subcategory(request,catid):
    subcategory=Subcategory.objects.filter(parent_cat_id=catid)
    print(subcategory)
    context={
        'subcategory':subcategory
    }
    return render(request,'adminapp/dropdown_e.html',context)

@login_required(login_url='adminlogin')
def add_product(request):
    if request.method=="POST":
        if request.POST.get('productname') and request.POST.get('price') and request.POST.get('productdesc') and request.POST.get('selectsubcategory') and request.FILES['productimage1'] and request.FILES['productimage2'] and request.FILES['productimage3'] and request.FILES['productimage4']:
        
            productname=request.POST.get('productname')
            if Product.objects.filter(product_name=productname).exists():
                messages.error(request,'Product Name Already Exists')
                return redirect('products')
            product=Product()
            product.product_name=productname
            product.slug=slugify(productname)
            product.price=request.POST.get('price')
            product.product_desc=request.POST.get('productdesc')
            product.subcategory_id=request.POST.get('selectsubcategory')
            product.category_id=request.POST.get('selectcategory')

            if 'productimage1' in request.FILES:
                product.img1=request.FILES['productimage1']
            if 'productimage2' in request.FILES:
                product.img2=request.FILES['productimage2']
            if 'productimage3' in request.FILES:
                product.img3=request.FILES['productimage3']
            if 'productimage4' in request.FILES:
                product.img4=request.FILES['productimage4']

            product.save()
            messages.success(request,"New Product Added!")
            return redirect('products')
        else:
            messages.error(request,'Required all Fields...')
            return redirect('addproduct')

    return redirect('addingproduct')
# have some problems with product edit

@login_required(login_url='adminlogin')
def edit_product(request, id):
    categories = Category.objects.all()
    subcategories = Subcategory.objects.all()
    product = get_object_or_404(Product, id=id)

    if request.method == 'POST':
        productname = request.POST.get('productname')
        if Product.objects.filter(product_name=productname).exclude(id=id).exists():
            messages.error(request, 'Product Name Already Exists')
            return redirect('products')

        product.product_name = productname
        product.slug = slugify(productname)
        product.price = request.POST.get('price')
        product.product_desc = request.POST.get('productdesc')

        subcategory_id = request.POST.get('selectsubcategory')
        category_id = request.POST.get('selectcategory')

        subcategory = get_object_or_404(Subcategory, id=subcategory_id)
        category = get_object_or_404(Category, id=category_id)

        product.subcategory = subcategory
        product.category = category

        if 'productimage1' in request.FILES:
            product.img1 = request.FILES['productimage1']
        if 'productimage2' in request.FILES:
            product.img2 = request.FILES['productimage2']
        if 'productimage3' in request.FILES:
            product.img3 = request.FILES['productimage3']
        if 'productimage4' in request.FILES:
            product.img4 = request.FILES['productimage4']

        try:
            product.save()
            messages.success(request, "Product Updated Successfully!")
        except Exception as e:
            messages.error(request, f"Error: {e}")

        return redirect('products')

    return render(request, 'adminapp/edit-product.html', {
        'product': product,
        'categories': categories,
        'subcategory': subcategories
    })

@login_required(login_url='adminlogin')
def delete_product(request,id):
    Product.objects.get(id=id).delete()
    return redirect('products')



# Product variation
# @login_required(login_url='adminlogin')
# def add_size(request):
#     if request.method=="POST":
#         if request.POST.get('selectcolor') and request.POST.get('size'):
#             seccolor=request.POST.get('selectcolor')
#             size=request.POST.get('size')
#             if Size.objects.filter(color=seccolor,size_value=size).exists():
#                 messages.error(request,"already exists")
#                 return redirect('variations')
#             else:
#                 s=Size()
#                 s.color_id=seccolor
#                 s.size_value=size
#                 s.save()
#                 messages.success(request,'added new size')
#                 return redirect('variations')
#         else:
#             messages.error(request,'Required all fields!!')
#             return redirect('variations')


@login_required(login_url='adminlogin')
def add_color(request):
    if request.method=='POST':
        if request.POST['color']:
            color=request.POST.get('color')
            if Color.objects.filter(color_value=color).exists():
                messages.error(request,"Already exists")
                return redirect('variations')
            else:
                c=Color()
                c.color_value=color
                c.save()
                messages.success(request,"added new color")
                return redirect('products')
        else:
            messages.error(request,"Required all fields!!")
            return redirect('variations')

@login_required(login_url='adminlogin')
def variations(request):
    color=Color.objects.all()
    return render(request,'adminapp/variations.html',{
        'color':color
        })

@login_required(login_url='adminlogin')
def add_variations(request, id):
    product = Product.objects.get(id=id)
    colors = Color.objects.all()

    if request.method == "POST":
        color_id = request.POST.get('selectcolor')
        color = Color.objects.get(id=color_id)

        stock = request.POST.get('stock')
        img1 = request.FILES.get('img1')
        img2 = request.FILES.get('img2')
        img3 = request.FILES.get('img3')
        img4 = request.FILES.get('img4')

        variation = Variations()
        variation.product = product
        variation.color = color
        variation.stock = stock
        variation.img1 = img1
        variation.img2 = img2
        variation.img3 = img3
        variation.img4 = img4
        variation.save()
        print("data saved ")

        messages.success(request, "Variation added.")
        return redirect('add_variations', id=id)
        # return redirect(reverse('products', kwargs={'id': id}))
    
    variations = product.variations.all()
    
    return render(request, 'adminapp/add_variations.html', {
        'colors': colors,
        'product': product,
        'variations': variations,
    })

@login_required(login_url='adminlogin')
def product_list(request):
    products = Product.objects.all()
    return render(request, 'adminapp/product_list.html', {'products': products})


# Order Management
@login_required(login_url='adminlogin')
def order_items(request):
    order_items=OrderItem.objects.order_by("-id").all()
    paginator=Paginator(order_items,10)
    page=request.GET.get('page')
    paged_orders_list=paginator.get_page(page) 
    
    
    
    # Calculate the coupon_price for each order_item
    for total in order_items:
        totalprice = total.price * total.quantity
       
    context={
        
        'order_items':paged_orders_list,
        'totalprice':totalprice,
    }
    return render(request,'adminapp/orderitems.html',context)

@login_required(login_url='adminlogin')
def order_details(request, order_id):
    try:
        order = Order.objects.get(id=order_id)
    except Order.DoesNotExist:
        return HttpResponse("Order not found")

    context = {
        'order': order
        
    }
    return render(request, 'adminapp/order_details.html', context)


# Offer Managment
@login_required(login_url='adminlogin')
def category_offer(request):
    cat_offer=CategoryOffer.objects.all()
    paginator=Paginator(cat_offer,10)
    page=request.GET.get('page')
    paged_category_offer_list=paginator.get_page(page)
    return render(request,'adminapp/category_offer.html',{
        'cat_offer':paged_category_offer_list
    })



@login_required(login_url='adminlogin')
def add_category_offer(request):
    form=CategoryOfferForm()
    if request.method=='POST':
        form=CategoryOfferForm(request.POST)
        discount=form.data['discount']
        if int(discount) <= 70:
            if form.is_valid():
                form.save()
                return redirect('category_offer')
            else:
                messages.error(request,'Already Exists!!')
                return redirect('add_category_offer')
        else:
            messages.error(request,'Percentage should be less than or equal to 70')
            return redirect('add_category_offer')
    context={
        'form':form
    }
    return render(request,'adminapp/add_category_offer.html',context)


@login_required(login_url='adminlogin')
def edit_category_offer(request,id):
    category_offer=get_object_or_404(CategoryOffer,id=id)
    form=CategoryOfferForm(request.POST or None,instance=category_offer)
    if form.is_valid():
        form.save()
        return redirect('category_offer')
    else:
        messages.error(request,"Please fill all fields *(discount should less than 70)")
    return render(request,'adminapp/edit_category_offer.html',{
        'form':form
    })

@login_required(login_url='adminlogin')
def delete_category_offer(request,id):
    cat_offer=get_object_or_404(CategoryOffer,id=id)
    cat_offer.delete()
    messages.error(request,'Deleted')
    return redirect('category_offer')

@login_required(login_url='adminlogin')
def subcategory_offer(request):
    sub_offers=SubcategoryOffer.objects.all()
    paginator=Paginator(sub_offers,10)
    page=request.GET.get('page')
    paged_sub_offer_list=paginator.get_page(page)
    return render(request,'adminapp/subcategory_offer.html',{
        'sub_offer':paged_sub_offer_list
    })


@login_required(login_url='adminlogin')
def add_subcategory_offer(request):
    form=SubcategoryOfferForm()
    if request.method=='POST':
        form=SubcategoryOfferForm(request.POST)
        discount=form.data['discount']
        if int(discount) <= 70:
            form=SubcategoryOfferForm(request.POST)
            if form.is_valid():
                form.save()
                return redirect('subcategory_offer')
            else:
                messages.error(request,'Already Exists!!')
                return redirect('add_subcategory_offer')
        else:
            messages.error(request,'Percentage should be less than or equal to 70')
            return redirect('add_subcategory_offer')   
    context={
        'form':form
    }
    return render(request,'adminapp/add_subcategory_offer.html',context)


@login_required(login_url='adminlogin')
def edit_subcategory_offer(request,id):
    subcategory_offer=get_object_or_404(SubcategoryOffer,id=id)
    form=SubcategoryOfferForm(request.POST or None,instance=subcategory_offer)
    if form.is_valid():
        form.save()
        return redirect('subcategory_offer')
    else:
        messages.error(request,'Percentage should be less than or equal to 70')
        
    return render(request,'adminapp/edit_subcategory_offer.html',{
        'form':form
    })

@login_required(login_url='adminlogin')
def delete_subcategory_offer(request,id):
    sub_offer=get_object_or_404(SubcategoryOffer,id=id)
    sub_offer.delete()
    messages.error(request,'Deleted')
    return redirect('subcategory_offer')


@login_required(login_url='adminlogin')
def product_offer(request):
    product_offers=ProductOffer.objects.all()
    paginator=Paginator(product_offers,10)
    page=request.GET.get('page')
    paged_product_offer_list=paginator.get_page(page)
    return render(request,'adminapp/product_offer.html',{
        'product_offer':paged_product_offer_list
    })


@login_required(login_url='adminlogin')
def add_product_offer(request):
    form=ProductOfferForm()
    if request.method=='POST':
        form=ProductOfferForm(request.POST)
        discount=form.data['discount']
        if int(discount) <= 70:
            form=ProductOfferForm(request.POST)
            if form.is_valid():
                form.save()
                return redirect('product_offer')
            else:
                messages.error(request,'Already Exists!!')
                return redirect('add_product_offer')
        else:
            messages.error(request,'Percentage should be less than or equal to 70')
            return redirect('add_product_offer')
    context={
        'form':form
    }
    return render(request,'adminapp/add_product_offer.html',context)


@login_required(login_url='adminlogin')
def edit_product_offer(request,id):
    product_offer=get_object_or_404(ProductOffer,id=id)
    form=ProductOfferForm(request.POST or None,instance=product_offer)
    if form.is_valid():
        form.save()
        return redirect('product_offer')
    else:
        messages.error(request,'Percentage should be less than or equal to 70')
    return render(request,'adminapp/edit_product_offer.html',{
        'form':form
    })


@login_required(login_url='adminlogin')
def delete_product_offer(request,id):
    pro_offer=get_object_or_404(ProductOffer,id=id)
    pro_offer.delete()
    messages.error(request,'Deleted')
    return redirect('product_offer')


@login_required(login_url='adminlogin')
def coupons(request):
    coupons_list=Coupon.objects.all()
    paginator=Paginator(coupons_list,10)
    page=request.GET.get('page')
    paged_coupons_list=paginator.get_page(page)
    return render(request,'adminapp/coupons.html',{
        'coupons':paged_coupons_list
    })


@login_required(login_url='adminlogin')
def add_coupons(request):
    form=CouponForm()
    if request.method=='POST':
        form=CouponForm(request.POST)
        discount=form.data['discount']
        if int(discount) <= 70:
            if form.is_valid():
                form.save()
                return redirect('coupons')
            else:
                messages.error(request,'Already Exists!!')
                return redirect('add_coupons')
        else:
            messages.error(request,'Percentage should be less than or equal to 70')
            return redirect('add_coupons')
    return render(request,'adminapp/add_coupons.html',{
        'form':form
    })


@login_required(login_url='adminlogin')
def edit_coupons(request,id):
    coupon=get_object_or_404(Coupon,id=id)
    form=CouponForm(request.POST or None,instance=coupon)
    if form.is_valid():
        form.save()
        return redirect('coupons')
    else:
        messages.error(request,'Percentage should be less than or equal to 70')
    return render(request,'adminapp/edit_coupon.html',{
        'form':form
    })

@login_required(login_url='adminlogin')
def delete_coupons(request,id):
    coupon=get_object_or_404(Coupon,id=id)
    coupon.delete()
    messages.error(request,'Deleted')
    return redirect('coupons')

# Report
@login_required(login_url='adminlogin')
def product_report(request):
    products=Product.objects.all()
    if request.GET.get('from'):
        product_date_from=datetime.datetime.strptime(request.GET.get('from'),"%Y-%m-%d")
        product_date_to=datetime.datetime.strptime(request.GET.get('to'),"%Y-%m-%d")

        product_date_to+=datetime.timedelta(days=1)
        products=Product.objects.filter(added_date__range=[product_date_from,product_date_to])    
    return render(request,'adminapp/product_report.html',{
        'products':products,
        
        
    })


class generateProductPdf(View):
    def get(self,request,*args,**kwargs):
        try:
            products=Product.objects.all()
            
        except:
            return HttpResponse("505 not found")
        data={
            'products':products
        }
        pdf=render_to_pdf('adminapp/product_pdf.html',data)
        return HttpResponse(pdf,content_type='application/pdf')



def render_to_pdf(template_src,context_dict={}):
    template=get_template(template_src)
    html=template.render(context_dict)
    result=BytesIO()
    pdf=pisa.pisaDocument(BytesIO(html.encode("ISO-8859-1")),result)
    if not pdf.err:
        return HttpResponse(result.getvalue(),content_type='application/pdf')
    return None

def product_csv(request):
    response=HttpResponse(content_type='text/csv')
    response[
        "Content-Disposition"
    ] = "attachement; filename=Product_Report.csv"

    writer=csv.writer(response)
    writer.writerow(
        [
            "Product Name",
            "Category Name",
            "Subcategory Name",
            "Price"
            "Stock"
        ]
    )
    products=Product.objects.all().order_by('id')
    for p in products:
        writer.writerow(
            [
                p.product_name,
                p.category.cat_name,
                p.subcategory.sub_name,
                p.price,
                p.stock,
            ]
        )
    return response



def product_excel(request):
    response=HttpResponse(content_type='application/ms-excel')
    response[
        "Content-Disposition"
    ] = "attachement; filename=Product_Report.xls"
    wb=xlwt.Workbook(encoding='utf-8')
    ws=wb.add_sheet('Product_Data')
    row_num=0
    font_style=xlwt.XFStyle()
    font_style.font.bold=True

    columns=[
            "Product Name",
            "Category Name",
            "Subcategory Name",
            "Price"
            "Stock"
    ]

    for col_num in range(len(columns)):
        ws.write(row_num,col_num,columns[col_num],font_style)

    font_style=xlwt.XFStyle()
    rows=Product.objects.all().values_list('product_name','category','subcategory','price','stock')

    for row in rows:
        row_num+=1

        for col_num in range(len(row)):
            ws.write(row_num,col_num,str(row[col_num]),font_style)

    wb.save(response)
    return response



# @login_required(login_url='adminlogin')
# def salesReport(request):
#     orders=Order.objects.all()
#     new_order_list=[]

#     for i in orders:
#         order_items=OrderItem.objects.filter(order_id=i.id)
#         for j in order_items:
#             item={
#                 'id':i.id,
#                 'ordered_date':i.created_at,
#                 'user':i.user,
#                 'price':j.price,
#                 'method':i.payment_mode,
#                 'status':j.status,


#             }
#             new_order_list.append(item)
#     paginator=Paginator(new_order_list,10)
#     page=request.GET.get('page')
#     paged_orders_list=paginator.get_page(page)    
#     return render(request,'adminapp/salesReport.html',{
#         'order':paged_orders_list,
#     })


# @login_required(login_url='adminlogin')
# def by_date(request):
#     from_date_str = request.GET.get('from')
#     to_date_str = request.GET.get('to')

#     if from_date_str and to_date_str:
#         sales_date_from = datetime.datetime.strptime(from_date_str, "%Y-%m-%d")
#         sales_date_to = datetime.datetime.strptime(to_date_str, "%Y-%m-%d")

#         if sales_date_from > sales_date_to:
#             messages.error(request, 'Invalid dates: "From" date cannot be later than "To" date.')
#             return redirect('salesReport')

#         sales_date_to += datetime.timedelta(days=1)
#         orders = Order.objects.filter(created_at__range=[sales_date_from, sales_date_to])

#         new_order_list = []

#         for i in orders:
#             order_items = OrderItem.objects.filter(order_id=i.id)
#             for j in order_items:
#                 item = {
#                     'id': i.id,
#                     'ordered_date': i.created_at,
#                     'user': i.user,
#                     'price': j.price,
#                     'method': i.payment_mode,
#                     'status': j.status,
#                 }
#                 new_order_list.append(item)
#     else:
#         messages.error(request, 'Select fields before submitting..!! ')
#         return redirect('salesReport')

#     return render(request, 'adminapp/salesReport.html', {
#         'order': new_order_list,
#     })

from django.utils import timezone

@login_required(login_url='adminlogin')
def salesReport(request):
    orders = Order.objects.all().order_by('-created_at')  # Sort orders by the most recent first
    new_order_list = []

    for order in orders:
        order_items = OrderItem.objects.filter(order_id=order.id)
        for order_item in order_items:
            item = {
                'id': order.id,
                'ordered_date': order.created_at,
                'user': order.user,
                'price': order_item.price,
                'method': order.payment_mode,
                'status': order_item.status,
            }
            new_order_list.append(item)

    paginator = Paginator(new_order_list, 10)
    page = request.GET.get('page')
    paged_orders_list = paginator.get_page(page)

    return render(request, 'adminapp/salesReport.html', {
        'order': paged_orders_list,
    })

@login_required(login_url='adminlogin')
def by_date(request):
    from_date_str = request.GET.get('from')
    to_date_str = request.GET.get('to')

    if from_date_str and to_date_str:
        sales_date_from = datetime.datetime.strptime(from_date_str, "%Y-%m-%d")
        sales_date_to = datetime.datetime.strptime(to_date_str, "%Y-%m-%d")

        if sales_date_from > sales_date_to:
            messages.error(request, 'Invalid dates: "From" date cannot be later than "To" date.')
            return redirect('salesReport')

        sales_date_to += datetime.timedelta(days=1)
        orders = Order.objects.filter(created_at__range=[sales_date_from, sales_date_to]).order_by('-created_at')

        new_order_list = []

        for order in orders:
            order_items = OrderItem.objects.filter(order_id=order.id)
            for order_item in order_items:
                item = {
                    'id': order.id,
                    'ordered_date': order.created_at,
                    'user': order.user,
                    'price': order_item.price,
                    'method': order.payment_mode,
                    'status': order_item.status,
                }
                new_order_list.append(item)
    else:
        messages.error(request, 'Select fields before submitting..!! ')
        return redirect('salesReport')

    return render(request, 'adminapp/salesReport.html', {
        'order': new_order_list,
    })





# def by_date(request):
#     if request.GET.get('from'):
#         sales_date_from=datetime.datetime.strptime(request.GET.get('from'),"%Y-%m-%d")
#         sales_date_to=datetime.datetime.strptime(request.GET.get('to'),"%Y-%m-%d")

#         sales_date_to+=datetime.timedelta(days=1)
#         orders=Order.objects.filter(created_at__range=[sales_date_from,sales_date_to])

#         new_order_list=[]

#         for i in orders:
#             order_items=OrderItem.objects.filter(order_id=i.id)
#             for j in order_items:
#                 item={
#                     'id':i.id,
#                     'ordered_date':i.created_at,
#                     'user':i.user,
#                     'price':j.price,
#                     'method':i.payment_mode,
#                     'status':j.status,


#                 }
#                 new_order_list.append(item)
#     else:
#         messages.error(request,'Select fields before submitting..!! ')
#         return redirect('salesReport')
        
#     return render(request,'adminapp/salesReport.html',{
#         'order':new_order_list,
#     })

class generatesalesReportPdf(View):
    def get(self,request,*args,**kwargs):
        try:
            orders=Order.objects.all()
            new_order_list=[]

            for i in orders:
                order_items=OrderItem.objects.filter(order_id=i.id)
                for j in order_items:
                    item={
                        'id':i.id,
                        'ordered_date':i.created_at,
                        'user':i.user,
                        'price':j.price,
                        'method':i.payment_mode,
                        'status':j.status,


                    }
                    new_order_list.append(item)
                    
        except:
            return HttpResponse("505 not found")
        data={
            'order':new_order_list
        }
        pdf=render_to_pdf('adminapp/salesReport_pdf.html',data)
        return HttpResponse(pdf,content_type='application/pdf')


@login_required(login_url='adminlogin')
def by_month(request):
    month=request.GET.get('month')
    orders=Order.objects.filter(created_at__month=month)
    new_order_list=[]

    for i in orders:
        order_items=OrderItem.objects.filter(order_id=i.id)
        for j in order_items:
            item={
                'id':i.id,
                'ordered_date':i.created_at,
                'user':i.user,
                'price':j.price,
                'method':i.payment_mode,
                'status':j.status,


            }
            new_order_list.append(item)
        
    return render(request,'adminapp/salesReport.html',{
        'order':new_order_list,
    })


@login_required(login_url='adminlogin')
def by_year(request):
    year=request.GET.get('year')
    orders=Order.objects.filter(created_at__year=year)
    new_order_list=[]

    for i in orders:
        order_items=OrderItem.objects.filter(order_id=i.id)
        for j in order_items:
            item={
                'id':i.id,
                'ordered_date':i.created_at,
                'user':i.user,
                'price':j.price,
                'method':i.payment_mode,
                'status':j.status,


            }
            new_order_list.append(item)
        
    return render(request,'adminapp/salesReport.html',{
        'order':new_order_list,
    })



def download_docx(request):
    document = Document()
    document.add_heading('Sales Report', 0)
    records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam')
    )
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=salesReport.docx'
    document.save(response)

    return response

from django.shortcuts import get_object_or_404

class GenerateInvoice(View):
    def get(self, request, id, *args, **kwargs):
        orders = get_object_or_404(Order, id=id, user=request.user)

        data = {
            'order_id': orders.id,
            'date': str(orders.created_at),
            'name': orders.user.first_name,
            'address': orders.address.address,
            'total_price': orders.total_price,
            'transaction_id': orders.payment_id,
            'payment_mode': orders.payment_mode,
            'user_email': orders.user.email,
            'orders': orders
        }

        pdf = render_to_pdf('invoice.html', data)
        if pdf:
            response = HttpResponse(pdf, content_type='application/pdf')
            response['Content-Disposition'] = 'filename="invoice.pdf"'
            return response

        return HttpResponse("Error generating PDF")




